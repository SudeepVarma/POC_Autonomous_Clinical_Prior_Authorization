"""
DocxLoader loads Word (.docx) documents and extracts their text and metadata.

Created on: 2026-08-05
Author: Sudeep Varma K
"""

from __future__ import annotations
from pathlib import Path
from docx import Document as WordDocument

from app.document.metadata import MetadataExtractor
from app.models.document import Document
from app.document.loaders.base import BaseLoader


class DocxLoader(BaseLoader):

    def load(
        self,
        path: Path,
    ) -> Document:

        doc = WordDocument(path)

        paragraphs = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        text = "\n".join(paragraphs)

        metadata = MetadataExtractor().from_file(
            path,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return Document(
            path=path,
            text=text,
            metadata=metadata,
        )